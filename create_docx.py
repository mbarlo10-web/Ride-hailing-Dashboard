#!/usr/bin/env python3
"""
Create Word document for Mini Project 3
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Mini Project 3 - Mark Barlow', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ============================================================================
# Q1: Justify your new design features
# ============================================================================
doc.add_heading('Q1: Justify your new design features (Why did you choose this layout/logic?)', 1)

doc.add_heading('Feature 1: Real-Time Parking Zone Status Grid', 2)

p = doc.add_paragraph()
p.add_run('Design Choice: ').bold = True
p.add_run('I chose to implement an 8x6 grid (48 zones) displaying parking/waiting area status with color-coded indicators. Each zone shows its identifier (e.g., "Zone A1", "Zone B2") and current occupancy count.')

doc.add_paragraph('Justification:')

items = [
    ('Visual Clarity', 'The grid layout provides an immediate, at-a-glance view of the entire waiting area. Passengers can quickly identify which zones are available without reading detailed text.'),
    ('Color-Coded System', 'Green (0 rides) = Available, Yellow (1-2 rides) = Moderate occupancy, Red (3+ rides) = Busy. This universal color language is intuitive and accessible to passengers of all languages and technical abilities, similar to traffic lights or airport signage.'),
    ('Spatial Mapping', 'The zones are mapped from actual spatial coordinates (x, y) in the dataset, creating a realistic representation of the physical layout. This helps passengers navigate to actual locations in the airport.'),
    ('Real-World Inspiration', 'This design was inspired by the demo video showing a ride-hailing dashboard at a train station in China, which demonstrated the effectiveness of grid-based zone displays for high-traffic transportation hubs.'),
    ('Decision Support', 'By showing occupancy levels, passengers can make informed decisions about where to wait, potentially reducing congestion in popular zones and improving overall flow.')
]

for item_title, item_text in items:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Feature 2: Active Ride Queue Display', 2)

p = doc.add_paragraph()
p.add_run('Design Choice: ').bold = True
p.add_run('A vertical list displaying active rides with license plate numbers, service types (Uber, Lyft, etc.), estimated wait times, and timestamps. The list updates in real-time and shows up to 20 most recent rides.')

doc.add_paragraph('Justification:')

items2 = [
    ('License Plate Identification', 'The most critical information for passengers is identifying their specific ride. Displaying license plates prominently (in monospace font for readability) allows passengers to quickly spot their vehicle.'),
    ('Service Type Display', 'Showing whether a ride is Uber, Lyft, or another service helps passengers filter information relevant to them, especially in areas where multiple services operate.'),
    ('Wait Time Estimates', 'Providing estimated wait times helps passengers plan their time and reduces anxiety about when their ride will arrive. This transparency improves the passenger experience.'),
    ('Chronological Ordering', 'Displaying rides sorted by most recent first ensures passengers see the latest information first, which is most relevant for those actively waiting.'),
    ('Real-Time Updates', 'The 5-second refresh rate provides a "live" feel without being overwhelming, striking a balance between information currency and system performance.'),
    ('Complementary to Zone Display', 'While the zone grid shows WHERE to wait, the ride queue shows WHAT is coming, creating a comprehensive information system that addresses both spatial and temporal needs.')
]

for item_title, item_text in items2:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Overall Layout Logic', 2)

p = doc.add_paragraph()
p.add_run('Two-Panel Design: ').bold = True
p.add_run('Left Panel (Parking Zones): Spatial information - WHERE to go. Right Panel (Active Rides): Temporal information - WHAT is coming. This split-screen approach separates concerns logically and allows passengers to focus on either spatial navigation or ride tracking as needed.')

p = doc.add_paragraph()
p.add_run('Airport-Style Aesthetic: ').bold = True
p.add_run('Dark blue gradient background matches airport display standards, high contrast for visibility on large screens, large fonts for readability from a distance, real-time clock for time awareness, and QR code for additional information access. This design ensures the dashboard fits seamlessly into an airport environment while providing maximum utility to passengers.')

doc.add_page_break()

# ============================================================================
# Q2: Describe your experience
# ============================================================================
doc.add_heading('Q2: Describe your experience of Mini Project 3. Please reflect on your experience of using Cursor and GitHub during your implementation.', 1)

doc.add_heading('Overall Experience', 2)
doc.add_paragraph('Mini Project 3 was an excellent opportunity to transform a data analysis project into a real-world application. The transition from static visualizations to an interactive, real-time dashboard required thinking about user experience, system architecture, and practical deployment considerations.')

doc.add_heading('Experience with Cursor', 2)

doc.add_paragraph('Positive Aspects:')

cursor_pros = [
    ('AI-Powered Assistance', 'Cursor\'s AI coding assistant was invaluable for quickly implementing complex features. When building the Flask web application, the AI helped generate boilerplate code, suggest best practices, and debug issues efficiently.'),
    ('Context Awareness', 'Cursor\'s ability to understand the entire codebase context made it particularly useful. When I needed to update multiple files (like adding QR code functionality), Cursor could see how changes in dashboard.py affected dashboard.html and suggest appropriate updates.'),
    ('Rapid Iteration', 'The ability to ask questions and get immediate code suggestions accelerated development significantly. For example, when implementing the QR code with dynamic IP addresses, Cursor helped identify the networking issue and provide a solution quickly.'),
    ('Error Debugging', 'When encountering issues like port conflicts or import errors, Cursor could analyze error messages and suggest fixes, saving considerable debugging time.'),
    ('Code Quality', 'The AI assistant helped maintain consistent code style and suggested improvements, such as adding error handling and making the QR code URL dynamic.')
]

for item_title, item_text in cursor_pros:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_paragraph('Challenges:')

cursor_cons = [
    ('Learning Curve', 'Initially, understanding how to effectively prompt Cursor required some trial and error. Learning to be specific about requirements (like "create a Flask route" vs. "make a web page") improved results.'),
    ('Over-Reliance Risk', 'There were moments where I had to step back and understand the code myself rather than just accepting AI suggestions, which is important for learning.'),
    ('Context Limitations', 'Sometimes Cursor would suggest solutions that didn\'t perfectly fit the existing codebase structure, requiring manual adjustments.')
]

for item_title, item_text in cursor_cons:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Experience with GitHub', 2)

doc.add_paragraph('Positive Aspects:')

github_pros = [
    ('Version Control', 'GitHub provided a safety net for experimentation. Knowing I could revert changes gave confidence to try new features and refactor code.'),
    ('Documentation', 'Using GitHub\'s README and markdown files helped document the project structure and features, making it easier to understand the codebase later.'),
    ('Collaboration Potential', 'While working solo, GitHub\'s structure made it clear how the project could be shared or collaborated on in the future.'),
    ('Project Organization', 'The repository structure (src/, templates/, assets/, etc.) helped organize the project logically and made navigation easier.')
]

for item_title, item_text in github_pros:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_paragraph('Challenges:')

github_cons = [
    ('Commit Discipline', 'Initially, I didn\'t commit as frequently as I should have. Learning to make smaller, more frequent commits would have been better practice.'),
    ('Branch Management', 'For a solo project, branching felt unnecessary, but I recognize it would be valuable in team settings.'),
    ('Documentation Balance', 'Finding the right amount of documentation (not too little, not too much) required some iteration.')
]

for item_title, item_text in github_cons:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Reflection on the Process', 2)
doc.add_paragraph('The combination of Cursor and GitHub created a powerful development environment: Cursor accelerated coding and problem-solving, while GitHub provided structure, safety, and documentation. This project demonstrated how modern AI-assisted development tools can enhance productivity while still requiring human judgment and understanding. The experience highlighted the importance of understanding the code you\'re writing, even when AI helps generate it.')

doc.add_page_break()

# ============================================================================
# Q3: Weigh pros and cons
# ============================================================================
doc.add_heading('Q3: Weigh the pros and cons of using Cursor & GitHub for your work projects.', 1)

doc.add_heading('Cursor - Pros and Cons', 2)

doc.add_heading('Pros:', 3)

cursor_pros_detailed = [
    ('Productivity Boost', 'Significantly faster code generation and implementation. Reduces time spent on boilerplate and repetitive tasks. Can help with unfamiliar technologies or frameworks.'),
    ('Learning Accelerator', 'Exposes you to new patterns and best practices. Provides explanations alongside code suggestions. Helps understand complex concepts through examples.'),
    ('Error Resolution', 'Quickly identifies and suggests fixes for bugs. Explains error messages in plain language. Can debug across multiple files simultaneously.'),
    ('Code Quality', 'Suggests improvements and optimizations. Helps maintain consistent coding style. Can identify potential security issues.'),
    ('Documentation Assistance', 'Generates docstrings and comments. Helps create README files and documentation. Can explain code functionality.'),
    ('Context Awareness', 'Understands your entire codebase. Maintains consistency across files. Remembers previous conversation context.')
]

for item_title, item_text in cursor_pros_detailed:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Cons:', 3)

cursor_cons_detailed = [
    ('Over-Dependence Risk', 'May reduce deep understanding of code. Could create gaps in fundamental knowledge. Risk of accepting code without comprehension.'),
    ('Quality Variability', 'Suggestions aren\'t always correct or optimal. May require significant debugging. Sometimes generates overly complex solutions.'),
    ('Learning Curve', 'Requires learning effective prompting techniques. Need to understand when to trust vs. verify suggestions. Can be time-consuming to refine prompts.'),
    ('Context Limitations', 'May not fully understand project-specific requirements. Can suggest generic solutions that don\'t fit perfectly. May miss subtle business logic requirements.'),
    ('Privacy Concerns', 'Code may be sent to external servers. Sensitive information could be exposed. Requires trust in the service provider.'),
    ('Cost', 'Premium features may require subscription. Could become expensive for teams. Free tier may have limitations.')
]

for item_title, item_text in cursor_cons_detailed:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('GitHub - Pros and Cons', 2)

doc.add_heading('Pros:', 3)

github_pros_detailed = [
    ('Version Control', 'Complete history of all changes. Ability to revert mistakes easily. Track evolution of codebase over time.'),
    ('Collaboration', 'Excellent for team projects. Pull requests enable code review. Issue tracking and project management.'),
    ('Backup and Safety', 'Cloud-based backup of code. Protection against local data loss. Can recover from any point in history.'),
    ('Documentation', 'README files for project overview. Wiki for detailed documentation. Markdown support for formatted docs.'),
    ('Integration', 'Works with CI/CD pipelines. Integrates with many development tools. Extensive ecosystem of integrations.'),
    ('Open Source Community', 'Access to millions of projects. Learn from others\' code. Contribute to open source projects.'),
    ('Professional Standard', 'Industry-standard tool. Expected skill in most tech jobs. Demonstrates professional practices.')
]

for item_title, item_text in github_pros_detailed:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Cons:', 3)

github_cons_detailed = [
    ('Learning Curve', 'Git commands can be complex. Merge conflicts can be confusing. Requires understanding of branching strategies.'),
    ('Overhead for Small Projects', 'May feel excessive for solo, simple projects. Requires discipline to commit regularly. Can slow down rapid prototyping.'),
    ('Public Repository Concerns', 'Default public repos may expose code. Need to be careful with sensitive data. Private repos require paid plan (for teams).'),
    ('Complexity', 'Advanced features can be overwhelming. Many options and configurations. Can be intimidating for beginners.'),
    ('Dependency', 'Reliance on external service. Requires internet connection. Service outages affect workflow.'),
    ('Storage Limitations', 'Large files require Git LFS. Repository size limits. May need to manage what\'s tracked.')
]

for item_title, item_text in github_cons_detailed:
    p = doc.add_paragraph(item_title + ': ', style='List Bullet')
    p.add_run(item_text)

doc.add_heading('Overall Assessment for Work Projects', 2)

p = doc.add_paragraph()
p.add_run('For Cursor:').bold = True
p.add_run(' Best for: Rapid prototyping, learning new technologies, debugging, code generation, documentation. Use with caution: Critical systems requiring deep understanding, security-sensitive code, complex business logic. Recommendation: Excellent tool for productivity, but always review and understand AI-generated code. Use as a powerful assistant, not a replacement for coding knowledge.')

p = doc.add_paragraph()
p.add_run('For GitHub:').bold = True
p.add_run(' Best for: Team collaboration, version control, project management, code review, professional development. Use with caution: Very small personal projects (may be overkill), projects with extremely large binary files. Recommendation: Essential tool for professional development. The benefits far outweigh the learning curve, and it\'s become a standard in the industry.')

p = doc.add_paragraph()
p.add_run('Combined Use:').bold = True
p.add_run(' The combination of Cursor and GitHub is particularly powerful: Cursor accelerates development, GitHub provides safety net and collaboration, together they create an efficient, professional workflow.')

p = doc.add_paragraph()
p.add_run('Conclusion:').bold = True
p.add_run(' Both tools are valuable additions to a developer\'s toolkit. Cursor enhances productivity and learning, while GitHub provides essential version control and collaboration capabilities. The key is using them appropriately - leveraging their strengths while maintaining understanding and control over your codebase.')

# Project Information
doc.add_page_break()
doc.add_heading('Project Information', 2)
doc.add_paragraph('Name: Mark Barlow')
doc.add_paragraph('Project: Mini Project 3 - Ride-Hailing Dashboard')
doc.add_paragraph('Date: December 2025')

# Save document
doc.save('MiniProject3_MarkBarlow.docx')
print("✓ Word document created: MiniProject3_MarkBarlow.docx")
print("  Location: /Users/markbarlow/mini-project-3-mark-barlow/MiniProject3_MarkBarlow.docx")

